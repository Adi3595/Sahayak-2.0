# backend/utils.py
import os
import uuid
import re
import logging
from typing import Optional, List, Dict
import docx
from PyPDF2 import PdfReader
import pytesseract
from PIL import Image
import io
from fpdf import FPDF
from bs4 import BeautifulSoup

# Set up logging
logger = logging.getLogger(__name__)


async def save_upload(file, upload_folder: str, allowed_extensions: set = None) -> str:
    """Save uploaded file with security checks"""
    if allowed_extensions is None:
        allowed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.jpg', '.jpeg', '.png'}
    
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    if file_ext not in allowed_extensions:
        raise ValueError(f"File type {file_ext} not allowed")
    
    # Create upload folder if it doesn't exist
    os.makedirs(upload_folder, exist_ok=True)
    
    filename = f"{uuid.uuid4()}{file_ext}"
    file_path = os.path.join(upload_folder, filename)
    
    try:
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
    except Exception as e:
        logger.error(f"Failed to save upload {filename}: {str(e)}")
        raise
    
    return file_path

def extract_text_from_file(file_path: str, page_start: int = 1, page_end: Optional[int] = None) -> str:
    """Extract text from file with page range support"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.pdf':
            return extract_text_from_pdf(file_path, page_start, page_end)
        elif file_ext in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return extract_text_from_txt(file_path)
        else:
            # Try OCR for images
            return extract_text_with_ocr(file_path)
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {str(e)}")
        raise Exception(f"Failed to extract text: {str(e)}")

def extract_text_from_pdf(file_path: str, page_start: int = 1, page_end: Optional[int] = None) -> str:
    """Extract text from PDF with page range"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            total_pages = len(reader.pages)
            
            # Validate page range
            actual_start = max(1, min(page_start, total_pages))
            if page_end:
                actual_end = min(page_end, total_pages)
            else:
                actual_end = total_pages
            
            if actual_start > actual_end:
                raise ValueError(f"Invalid page range: {actual_start} to {actual_end}")
            
            # Extract text from specified pages
            for page_num in range(actual_start - 1, actual_end):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                
                if page_text and page_text.strip():
                    text += page_text + "\n\n"
                else:
                    logger.info(f"Page {page_num + 1} appears to have no extractable text")
    
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {str(e)}")
        raise
    
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        return text
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
        raise

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        # If all encodings fail, use utf-8 with errors ignored
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except Exception as e:
        logger.error(f"TXT extraction failed for {file_path}: {str(e)}")
        raise

def extract_text_with_ocr(file_path: str) -> str:
    """Extract text from image using OCR"""
    try:
        image = Image.open(file_path)
        return pytesseract.image_to_string(image)
    except Exception as e:
        logger.error(f"OCR failed for {file_path}: {str(e)}")
        raise Exception(f"OCR failed: {str(e)}")

def extract_topic_from_text(text: str, max_length: int = 50) -> str:
    """Extract a meaningful topic from the source text"""
    try:
        # Use first few meaningful lines to determine topic
        lines = text.split('\n')
        meaningful_lines = [line.strip() for line in lines if line.strip() and len(line.strip()) > 10]
        
        if meaningful_lines:
            # Take first meaningful line and truncate
            first_line = meaningful_lines[0]
            # Remove special characters and extra spaces
            topic = re.sub(r'[^\w\s]', '', first_line)
            topic = ' '.join(topic.split()[:10])  # Take first 10 words
            topic = topic[:max_length].rsplit(' ', 1)[0]  # Don't cut words
            return topic + ('...' if len(first_line) > max_length else '')
        else:
            return "Educational Worksheet"
            
    except Exception as e:
        logger.warning(f"Topic extraction failed: {e}")
        return "Educational Worksheet"

def format_grades_display(grades: List[str]) -> str:
    """Format grades in a nice display format"""
    if not grades:
        return ""
    
    # Remove duplicates and sort
    unique_grades = sorted(set(grades), key=lambda x: int(x) if x.isdigit() else 0)
    
    if len(unique_grades) == 1:
        return f"Grade {unique_grades[0]}"
    elif len(unique_grades) == 2:
        return f"Grades {unique_grades[0]} and {unique_grades[1]}"
    else:
        # For multiple grades: "Grades 3, 4, and 5"
        return f"Grades {', '.join(unique_grades[:-1])}, and {unique_grades[-1]}"

def generate_worksheets_with_groq(
    source_text: str, 
    subject: str, 
    grades: List[str], 
    difficulty: str, 
    groq_client,
    question_types: List[str] = None,  # ADD THIS PARAMETER
    max_source_length: int = 6000
) -> Dict[str, str]:
    """Generate differentiated worksheets using Groq API with specified question types"""
    
    # Set default question types if none provided
    if question_types is None:
        question_types = ["mcq", "theory"]
    
    # Truncate source text intelligently
    if len(source_text) > max_source_length:
        source_text = source_text[:max_source_length].rsplit(' ', 1)[0] + "..."
    
    worksheets = {}
    
    for grade in grades:
        prompt = create_worksheet_prompt(source_text, subject, grade, difficulty, question_types)
        
        try:
            response = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {
                        "role": "system",
                        "content": """You are an expert educational content creator. Create engaging, 
                        age-appropriate worksheets that help students learn effectively. 
                        Always include the specific question types requested in the prompt.
                        Format your response in clean HTML with clear sections."""
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.7,
                max_tokens=4000,  # Increased for comprehensive worksheets
                top_p=0.9
            )
            
            worksheet_content = response.choices[0].message.content
            worksheets[grade] = worksheet_content
            
        except Exception as e:
            logger.error(f"Failed to generate worksheet for grade {grade}: {str(e)}")
            worksheets[grade] = create_error_worksheet(grade, str(e))
    
    return worksheets

def create_worksheet_prompt(source_text: str, subject: str, grade: str, difficulty: str, question_types: List[str]) -> str:
    """Create a structured prompt for worksheet generation with specific question types"""
    
    # Create question type descriptions
    question_sections = []
    
    if "mcq" in question_types:
        question_sections.append("""
MULTIPLE CHOICE QUESTIONS (8-10 questions):
- Create questions with 4 options each (A, B, C, D)
- Mark the correct answer clearly with (Correct)
- Cover key concepts from the source material
- Include a mix of factual and conceptual questions
        """)
    
    if "theory" in question_types:
        question_sections.append("""
THEORY/SHORT ANSWER QUESTIONS (5-7 questions):
- Create thought-provoking questions requiring 2-4 sentence answers
- Include questions that test comprehension and application
- Provide space for detailed responses
- Cover both factual recall and analytical thinking
        """)
    
    if "fill" in question_types:
        question_sections.append("""
FILL IN THE BLANKS (6-8 questions):
- Create sentences with key terms missing
- Include a word bank if appropriate for the grade level
- Focus on important vocabulary and concepts
        """)
    
    if "truefalse" in question_types:
        question_sections.append("""
TRUE/FALSE QUESTIONS (5-7 questions):
- Create clear, unambiguous statements
- Include explanations for false statements
- Cover common misconceptions
        """)
    
    question_types_text = "\n".join(question_sections)
    
    return f"""
Create a comprehensive educational worksheet for {subject} at {grade} grade level.

SOURCE CONTEXT:
{source_text}

DIFFICULTY LEVEL: {difficulty}
GRADE LEVEL: {grade}
QUESTION TYPES REQUIRED: {', '.join(question_types).upper()}

Worksheet Structure:

1. LEARNING OBJECTIVES: 3-5 clear, measurable objectives
2. ENGAGING INTRODUCTION: Brief overview connecting to real-world relevance
3. PRACTICE EXERCISES:

{question_types_text}

4. ANSWER KEY: Provide complete solutions for all questions including:
   - Correct answers for MCQs with brief explanations
   - Sample answers for theory questions
   - Completed fill-in-the-blanks
   - True/False answers with explanations

Differentiation Guidelines:
- Grade level appropriateness for {grade}
- {difficulty.capitalize()} difficulty level
- Clear, age-appropriate language
- Varied question formats to engage different learning styles

Format in clean HTML with:
- Clear headings (<h2> for main sections, <h3> for subsections)
- Proper paragraph tags (<p>)
- Numbered lists for questions (<ol>)
- Bullet points for options (<ul>)
- Bold text for important concepts (<strong>)
- Adequate spacing between sections

Make the worksheet visually appealing, easy to follow, and ready for classroom use.
"""

def create_error_worksheet(grade: str, error: str) -> str:
    """Create a fallback worksheet when generation fails"""
    return f"""
    <div style="font-family: Arial, sans-serif; padding: 20px; border: 2px solid #ff6b6b; border-radius: 10px; background: #fff5f5;">
        <h2 style="color: #e53e3e;">Worksheet for Grade {grade}</h2>
        <p style="color: #718096;">We encountered an issue generating this worksheet:</p>
        <p style="color: #e53e3e; font-weight: bold;">{error}</p>
        <p style="color: #718096;">Please try again or contact support if the problem persists.</p>
    </div>
    """

def create_formatted_docx(worksheet_content: str, subject: str, grade: str, difficulty: str, output_path: str, question_types: List[str] = None):
    """Create a well-formatted DOCX worksheet with question type support"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'{subject} Worksheet - Grade {grade}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle with difficulty and question types
        subtitle_text = f'Difficulty Level: {difficulty.title()}'
        if question_types:
            subtitle_text += f' | Question Types: {", ".join(question_types).title()}'
        subtitle = doc.add_paragraph(subtitle_text)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add spacing
        
        # Convert HTML content to formatted DOCX
        try:
            # Parse HTML content
            soup = BeautifulSoup(worksheet_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Process each element
            for element in soup.find_all(True):
                if element.name in ['h1', 'h2', 'h3']:
                    level = int(element.name[1])
                    heading_text = clean_text_for_pdf(element.get_text().strip())
                    if heading_text:
                        heading = doc.add_heading(heading_text, level=min(level-1, 2))
                        
                elif element.name == 'p':
                    text = clean_text_for_pdf(element.get_text().strip())
                    if text:
                        # Check if this is an MCQ option
                        if re.match(r'^[A-D][\.\)]', text.strip()):
                            p = doc.add_paragraph()
                            p.paragraph_format.left_indent = Inches(0.3)
                            p.add_run(text)
                        else:
                            doc.add_paragraph(text)
                        
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        li_text = clean_text_for_pdf(li.get_text().strip())
                        if li_text:
                            if element.name == 'ul':
                                p = doc.add_paragraph(li_text, style='List Bullet')
                            else:
                                p = doc.add_paragraph(li_text, style='List Number')
                            
                elif element.name == 'br':
                    doc.add_paragraph()
                    
        except Exception as e:
            # Fallback: simple text processing with better formatting
            logger.warning(f"HTML parsing failed, using fallback: {e}")
            clean_text = re.sub('<[^<]+?>', '', worksheet_content)
            clean_text = clean_text_for_pdf(clean_text)
            
            # Improved text processing for questions
            lines = clean_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                    
                # Detect and format questions
                if re.match(r'^\d+[\.\)]', line):  # Numbered questions
                    p = doc.add_paragraph()
                    p.add_run(line).bold = True
                elif re.match(r'^[A-D][\.\)]', line):  # MCQ options
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.add_run(line)
                elif any(keyword in line.lower() for keyword in ['answer:', 'correct:', 'explanation:']):
                    p = doc.add_paragraph()
                    p.add_run(line).bold = True
                else:
                    doc.add_paragraph(line)
        
        doc.save(output_path)
        logger.info(f"DOCX saved successfully: {output_path}")
        
    except Exception as e:
        logger.error(f"Failed to create DOCX: {str(e)}")
        raise Exception(f"DOCX creation failed: {str(e)}")

def create_formatted_pdf(worksheet_content: str, subject: str, grade: str, difficulty: str, output_path: str, question_types: List[str] = None):
    """Create a formatted PDF worksheet with Unicode support and question type formatting"""
    try:
        # Create PDF with Unicode support
        pdf = FPDF()
        
        # Add a Unicode-compatible font
        try:
            pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
            pdf.add_font('DejaVu', 'B', 'DejaVuSans-Bold.ttf', uni=True)
            pdf.add_font('DejaVu', 'I', 'DejaVuSans-Oblique.ttf', uni=True)
            font_family = 'DejaVu'
        except:
            font_family = 'Arial'
            logger.warning("DejaVu fonts not found, using Arial (limited Unicode support)")
        
        pdf.add_page()
        
        # Set font for title
        pdf.set_font(font_family, 'B', 16)
        title_text = clean_text_for_pdf(f'{subject} Worksheet - Grade {grade}')
        pdf.cell(0, 10, title_text, 0, 1, 'C')
        
        # Subtitle with question types
        subtitle_text = f'Difficulty Level: {difficulty.title()}'
        if question_types:
            subtitle_text += f' | Question Types: {", ".join(question_types).title()}'
        pdf.set_font(font_family, 'I', 12)
        pdf.cell(0, 10, clean_text_for_pdf(subtitle_text), 0, 1, 'C')
        pdf.ln(10)
        
        # Parse and add content with improved formatting
        try:
            soup = BeautifulSoup(worksheet_content, 'html.parser')
            
            for element in soup.find_all(True):
                text = clean_text_for_pdf(element.get_text().strip())
                if not text:
                    continue
                    
                if element.name in ['h1', 'h2', 'h3']:
                    # Main headings
                    pdf.set_font(font_family, 'B', 14)
                    pdf.multi_cell(0, 10, text)
                    pdf.set_font(font_family, size=12)
                    pdf.ln(5)
                elif element.name == 'p':
                    # Paragraphs - check for special formatting
                    if re.match(r'^[A-D][\.\)]', text):  # MCQ options
                        pdf.set_font(font_family, size=11)
                        pdf.cell(10)  # Indent
                        pdf.multi_cell(0, 8, text)
                    elif re.match(r'^\d+[\.\)]', text):  # Numbered questions
                        pdf.set_font(font_family, 'B', 12)
                        pdf.multi_cell(0, 8, text)
                        pdf.set_font(font_family, size=12)
                    else:
                        pdf.set_font(font_family, size=12)
                        pdf.multi_cell(0, 8, text)
                    pdf.ln(5)
                elif element.name in ['ul', 'ol']:
                    # Lists - handle MCQs and other lists
                    for li in element.find_all('li'):
                        li_text = clean_text_for_pdf(li.get_text().strip())
                        if element.name == 'ul':
                            # Bullet points for MCQs
                            pdf.set_font(font_family, size=11)
                            pdf.cell(10)
                            pdf.multi_cell(0, 8, f"• {li_text}")
                        else:
                            # Numbered lists
                            pdf.set_font(font_family, size=12)
                            pdf.multi_cell(0, 8, li_text)
                    pdf.ln(5)
                    
        except Exception as e:
            # Fallback: improved text processing
            logger.warning(f"HTML parsing failed for PDF, using fallback: {e}")
            clean_text = re.sub('<[^<]+?>', '', worksheet_content)
            clean_text = clean_text_for_pdf(clean_text)
            
            lines = clean_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    pdf.ln(5)
                    continue
                    
                # Improved formatting for different content types
                if re.match(r'^\d+[\.\)]', line):  # Questions
                    pdf.set_font(font_family, 'B', 12)
                    pdf.multi_cell(0, 8, line)
                    pdf.set_font(font_family, size=12)
                elif re.match(r'^[A-D][\.\)]', line):  # MCQ options
                    pdf.set_font(font_family, size=11)
                    pdf.cell(10)
                    pdf.multi_cell(0, 8, line)
                    pdf.set_font(font_family, size=12)
                elif any(keyword in line.lower() for keyword in ['answer key', 'correct answer', 'explanation']):
                    pdf.set_font(font_family, 'B', 12)
                    pdf.multi_cell(0, 8, line)
                    pdf.set_font(font_family, size=12)
                else:
                    pdf.set_font(font_family, size=12)
                    pdf.multi_cell(0, 8, line)
                pdf.ln(5)
        
        pdf.output(output_path)
        logger.info(f"PDF saved successfully: {output_path}")
        
    except Exception as e:
        logger.error(f"PDF creation failed: {e}")
        # Try alternative PDF creation method
        try:
            create_pdf_fallback(worksheet_content, subject, grade, difficulty, output_path, question_types)
        except Exception as fallback_error:
            logger.error(f"PDF fallback also failed: {fallback_error}")
            raise Exception(f"Failed to create PDF: {str(e)}")

def clean_text_for_pdf(text: str) -> str:
    """Clean text for PDF compatibility by replacing problematic Unicode characters"""
    if not text:
        return ""
    
    # Replace common problematic Unicode characters
    replacements = {
        '\u2022': '-',      # Bullet
        '\u25cf': '-',      # Black circle
        '\u25cb': '-',      # White circle
        '\u2013': '-',      # En dash
        '\u2014': '-',      # Em dash
        '\u2018': "'",      # Left single quotation
        '\u2019': "'",      # Right single quotation
        '\u201c': '"',      # Left double quotation
        '\u201d': '"',      # Right double quotation
        '\u00a0': ' ',      # Non-breaking space
        '\u00b0': ' deg',   # Degree symbol
        '\u00f7': '/',      # Division
        '\u00d7': 'x',      # Multiplication
    }
    
    # Apply replacements
    for unicode_char, replacement in replacements.items():
        text = text.replace(unicode_char, replacement)
    
    # Remove any other non-ASCII characters that might cause issues
    text = text.encode('ascii', 'ignore').decode('ascii')
    
    return text

def create_pdf_fallback(worksheet_content: str, subject: str, grade: str, difficulty: str, output_path: str, question_types: List[str] = None):
    """Fallback PDF creation method using reportlab if available"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # Create PDF
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        # Set up fonts
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
            font_name = 'DejaVuSans'
        except:
            font_name = 'Helvetica'
        
        # Add title
        c.setFont(font_name, 16)
        c.drawString(50, height - 50, f"{subject} Worksheet - Grade {grade}")
        c.setFont(font_name, 12)
        
        # Add subtitle with question types
        subtitle = f"Difficulty Level: {difficulty.title()}"
        if question_types:
            subtitle += f" | Question Types: {', '.join(question_types).title()}"
        c.drawString(50, height - 70, subtitle)
        
        # Parse and add content
        y_position = height - 100
        line_height = 14
        
        try:
            soup = BeautifulSoup(worksheet_content, 'html.parser')
            text_elements = []
            
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'li']):
                text = element.get_text().strip()
                if text:
                    if element.name in ['h1', 'h2', 'h3']:
                        text_elements.append(('heading', text))
                    elif element.name == 'li':
                        text_elements.append(('list', text))
                    else:
                        # Check for question patterns
                        if re.match(r'^\d+[\.\)]', text):
                            text_elements.append(('question', text))
                        elif re.match(r'^[A-D][\.\)]', text):
                            text_elements.append(('mcq_option', text))
                        else:
                            text_elements.append(('paragraph', text))
            
            for elem_type, text in text_elements:
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
                
                if elem_type == 'heading':
                    c.setFont(font_name, 14)
                    c.drawString(50, y_position, text[:80])
                    y_position -= line_height * 1.5
                elif elem_type == 'question':
                    c.setFont(font_name, 12)
                    c.drawString(50, y_position, text)
                    y_position -= line_height
                elif elem_type == 'mcq_option':
                    c.setFont(font_name, 10)
                    c.drawString(60, y_position, text)  # Indent options
                    y_position -= line_height
                elif elem_type == 'list':
                    c.setFont(font_name, 10)
                    c.drawString(55, y_position, f"• {text}")
                    y_position -= line_height
                else:
                    c.setFont(font_name, 10)
                    # Simple text wrapping
                    words = text.split()
                    lines = []
                    current_line = []
                    
                    for word in words:
                        test_line = ' '.join(current_line + [word])
                        if len(test_line) < 80:
                            current_line.append(word)
                        else:
                            lines.append(' '.join(current_line))
                            current_line = [word]
                    
                    if current_line:
                        lines.append(' '.join(current_line))
                    
                    for line in lines:
                        if y_position < 50:
                            c.showPage()
                            y_position = height - 50
                        c.drawString(50, y_position, line)
                        y_position -= line_height
                
                y_position -= 5  # Add spacing between elements
        
        except Exception as e:
            # Fallback to simple text
            logger.warning(f"ReportLab HTML parsing failed: {e}")
            clean_text = re.sub('<[^<]+?>', '', worksheet_content)
            lines = clean_text.split('\n')
            
            for line in lines:
                if line.strip():
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50
                    c.setFont(font_name, 10)
                    c.drawString(50, y_position, line.strip()[:100])
                    y_position -= line_height
        
        c.save()
        logger.info(f"Fallback PDF saved successfully: {output_path}")
        
    except ImportError:
        logger.error("ReportLab not available for fallback PDF creation")
        raise Exception("PDF creation failed and ReportLab fallback not available")
    except Exception as e:
        logger.error(f"Fallback PDF creation also failed: {e}")
        raise

def create_docx_from_text(content: str, output_path: str):
    """Create DOCX file from text content (legacy function)"""
    try:
        doc = docx.Document()
        
        # Add content to document
        if content.startswith('<'):
            # Simple HTML to text conversion for DOCX
            clean_content = re.sub('<[^<]+?>', '', content)
            clean_content = clean_text_for_pdf(clean_content)
            for paragraph in clean_content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        else:
            clean_content = clean_text_for_pdf(content)
            for paragraph in clean_content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        
        doc.save(output_path)
        logger.info(f"Legacy DOCX saved: {output_path}")
        
    except Exception as e:
        logger.error(f"Legacy DOCX creation failed: {str(e)}")
        raise

def create_pdf_from_html_optional(html_content: str, output_path: str):
    """Create PDF from HTML content - simplified version with Unicode support"""
    try:
        pdf = FPDF()
        
        # Add Unicode font
        try:
            pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
            font_family = 'DejaVu'
        except:
            font_family = 'Arial'
        
        pdf.add_page()
        pdf.set_font(font_family, size=12)
        
        # Simple text extraction from HTML with Unicode cleaning
        clean_text = re.sub('<[^<]+?>', '', html_content)
        clean_text = clean_text_for_pdf(clean_text)
        
        # Add text to PDF
        for line in clean_text.split('\n'):
            if line.strip():
                pdf.multi_cell(0, 10, line.strip())
        
        pdf.output(output_path)
        logger.info(f"Simple PDF saved: {output_path}")
        
    except Exception as e:
        logger.error(f"Simple PDF creation failed: {str(e)}")
        # Fallback: create a simple text file
        try:
            with open(output_path.replace('.pdf', '.txt'), 'w', encoding='utf-8') as f:
                f.write(html_content)
            logger.info(f"Created text file instead: {output_path.replace('.pdf', '.txt')}")
        except:
            pass
        raise Exception(f"PDF creation failed, created text file instead: {str(e)}")

def clean_html_content(html_content: str) -> str:
    """Clean and sanitize HTML content for display"""
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove potentially harmful tags
        for tag in soup(['script', 'style', 'meta', 'link']):
            tag.decompose()
        
        # Clean up the HTML
        cleaned_html = str(soup)
        
        # Limit content length for preview
        if len(cleaned_html) > 1000:
            cleaned_html = cleaned_html[:1000] + "..."
            
        return cleaned_html
        
    except Exception as e:
        logger.warning(f"HTML cleaning failed: {e}")
        # Return plain text if cleaning fails
        plain_text = re.sub('<[^<]+?>', '', html_content)
        return plain_text[:500] + "..." if len(plain_text) > 500 else plain_text

def validate_grade_levels(grades: List[str]) -> List[str]:
    """Validate and normalize grade levels"""
    valid_grades = ['k', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12']
    
    normalized_grades = []
    for grade in grades:
        grade = grade.lower().strip()
        if grade in valid_grades:
            normalized_grades.append(grade)
        else:
            logger.warning(f"Invalid grade level skipped: {grade}")
    
    return normalized_grades

def get_file_size_mb(file_path: str) -> float:
    """Get file size in MB"""
    try:
        size_bytes = os.path.getsize(file_path)
        return round(size_bytes / (1024 * 1024), 2)
    except:
        return 0.0

def cleanup_old_files(directory: str, max_age_hours: int = 24):
    """Clean up files older than specified hours"""
    try:
        import time
        current_time = time.time()
        
        if not os.path.exists(directory):
            return
        
        for filename in os.listdir(directory):
            file_path = os.path.join(directory, filename)
            if os.path.isfile(file_path):
                file_age = current_time - os.path.getctime(file_path)
                if file_age > max_age_hours * 3600:
                    os.remove(file_path)
                    logger.info(f"Cleaned up old file: {filename}")
                    
    except Exception as e:
        logger.error(f"Cleanup error for {directory}: {e}")

def format_question_types_display(question_types: List[str]) -> str:
    """Format question types for display"""
    if not question_types:
        return "Mixed Questions"
    
    type_display = {
        'mcq': 'Multiple Choice',
        'theory': 'Theory/Short Answer', 
        'fill': 'Fill in the Blanks',
        'truefalse': 'True/False'
    }
    
    display_names = [type_display.get(qtype, qtype.title()) for qtype in question_types]
    
    if len(display_names) == 1:
        return display_names[0]
    elif len(display_names) == 2:
        return f"{display_names[0]} and {display_names[1]}"
    else:
        return f"{', '.join(display_names[:-1])}, and {display_names[-1]}"